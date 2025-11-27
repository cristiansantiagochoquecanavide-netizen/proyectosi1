"""
Módulo para exportar datos clínicos a diferentes formatos
"""
from io import BytesIO
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from django.http import HttpResponse


def exportar_citas_a_excel(citas):
    """Exporta citas a archivo Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Citas"
    
    # Estilos
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Encabezados
    headers = ["ID", "Fecha y Hora", "Paciente", "Odontólogo", "Estado"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    
    # Datos
    for row, cita in enumerate(citas, 2):
        data = [
            cita.get('id_cita', ''),
            cita.get('fecha', ''),
            cita.get('id_paciente', {}).get('nombre', '') if isinstance(cita.get('id_paciente'), dict) else cita.get('id_paciente', ''),
            cita.get('id_odontologo', {}).get('nombre', '') if isinstance(cita.get('id_odontologo'), dict) else cita.get('id_odontologo', ''),
            cita.get('estado', ''),
        ]
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = border
            cell.alignment = Alignment(horizontal="left", vertical="center")
    
    # Ajustar ancho de columnas
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 15
    
    # Guardar en BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def exportar_atenciones_a_excel(atenciones):
    """Exporta atenciones a archivo Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Atenciones"
    
    # Estilos
    header_fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Encabezados
    headers = ["ID", "Paciente", "Odontólogo", "Fecha Inicio", "Fecha Fin", "Estado", "Observaciones"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    
    # Datos
    for row, atencion in enumerate(atenciones, 2):
        data = [
            atencion.get('id_atencion', ''),
            atencion.get('id_paciente', {}).get('nombre', '') if isinstance(atencion.get('id_paciente'), dict) else atencion.get('id_paciente', ''),
            atencion.get('id_odontologo', {}).get('nombre', '') if isinstance(atencion.get('id_odontologo'), dict) else atencion.get('id_odontologo', ''),
            atencion.get('fecha_inicio', ''),
            atencion.get('fecha_fin', ''),
            atencion.get('estado', ''),
            atencion.get('observaciones_generales', '')[:50] if atencion.get('observaciones_generales') else '',
        ]
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = border
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    # Ajustar ancho de columnas
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 20
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 30
    
    # Guardar en BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def exportar_citas_a_word(citas):
    """Exporta citas a documento Word"""
    doc = Document()
    
    # Título
    title = doc.add_heading('Reporte de Citas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha de generación
    fecha = doc.add_paragraph(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Espacio
    
    # Tabla
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    headers = ["ID", "Fecha y Hora", "Paciente", "Odontólogo", "Estado"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Formato header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Datos
    for cita in citas:
        row_cells = table.add_row().cells
        paciente = cita.get('id_paciente', {})
        if isinstance(paciente, dict):
            paciente = paciente.get('nombre', '')
        
        odontologo = cita.get('id_odontologo', {})
        if isinstance(odontologo, dict):
            odontologo = odontologo.get('nombre', '')
        
        row_cells[0].text = str(cita.get('id_cita', ''))
        row_cells[1].text = str(cita.get('fecha', ''))
        row_cells[2].text = str(paciente)
        row_cells[3].text = str(odontologo)
        row_cells[4].text = str(cita.get('estado', ''))
    
    # Estadísticas
    doc.add_paragraph()
    doc.add_heading('Estadísticas', level=2)
    doc.add_paragraph(f'Total de Citas: {len(citas)}')
    doc.add_paragraph(f'Confirmadas: {sum(1 for c in citas if c.get("estado") == "confirmada")}')
    doc.add_paragraph(f'Pendientes: {sum(1 for c in citas if c.get("estado") == "pendiente")}')
    doc.add_paragraph(f'Canceladas: {sum(1 for c in citas if c.get("estado") == "cancelada")}')
    
    # Guardar en BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def exportar_atenciones_a_word(atenciones):
    """Exporta atenciones a documento Word"""
    doc = Document()
    
    # Título
    title = doc.add_heading('Reporte de Atenciones', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha de generación
    fecha = doc.add_paragraph(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Espacio
    
    # Tabla
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    headers = ["ID", "Paciente", "Odontólogo", "Fecha Inicio", "Fecha Fin", "Estado", "Observaciones"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Datos
    for atencion in atenciones:
        row_cells = table.add_row().cells
        paciente = atencion.get('id_paciente', {})
        if isinstance(paciente, dict):
            paciente = paciente.get('nombre', '')
        
        odontologo = atencion.get('id_odontologo', {})
        if isinstance(odontologo, dict):
            odontologo = odontologo.get('nombre', '')
        
        row_cells[0].text = str(atencion.get('id_atencion', ''))
        row_cells[1].text = str(paciente)
        row_cells[2].text = str(odontologo)
        row_cells[3].text = str(atencion.get('fecha_inicio', ''))
        row_cells[4].text = str(atencion.get('fecha_fin', ''))
        row_cells[5].text = str(atencion.get('estado', ''))
        row_cells[6].text = str(atencion.get('observaciones_generales', '')[:100])
    
    # Estadísticas
    doc.add_paragraph()
    doc.add_heading('Estadísticas', level=2)
    doc.add_paragraph(f'Total de Atenciones: {len(atenciones)}')
    doc.add_paragraph(f'Finalizadas: {sum(1 for a in atenciones if a.get("estado") == "finalizada")}')
    doc.add_paragraph(f'En Curso: {sum(1 for a in atenciones if a.get("estado") == "en_curso")}')
    doc.add_paragraph(f'Canceladas: {sum(1 for a in atenciones if a.get("estado") == "cancelada")}')
    
    # Guardar en BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def exportar_citas_a_pdf(citas):
    """Exporta citas a PDF"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    elements = []
    
    # Estilos
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#4472C4'),
        spaceAfter=30,
        alignment=1  # CENTER
    )
    
    # Título
    title = Paragraph("Reporte de Citas", title_style)
    elements.append(title)
    
    # Fecha
    fecha_text = Paragraph(f'<b>Generado:</b> {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}', styles['Normal'])
    elements.append(fecha_text)
    elements.append(Spacer(1, 0.3*inch))
    
    # Tabla
    data = [["ID", "Fecha y Hora", "Paciente", "Odontólogo", "Estado"]]
    for cita in citas:
        paciente = cita.get('id_paciente', {})
        if isinstance(paciente, dict):
            paciente = paciente.get('nombre', '')
        
        odontologo = cita.get('id_odontologo', {})
        if isinstance(odontologo, dict):
            odontologo = odontologo.get('nombre', '')
        
        data.append([
            str(cita.get('id_cita', '')),
            str(cita.get('fecha', '')),
            str(paciente),
            str(odontologo),
            str(cita.get('estado', ''))
        ])
    
    table = Table(data, colWidths=[0.8*inch, 1.5*inch, 1.5*inch, 1.5*inch, 1.2*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Estadísticas
    stats_title = Paragraph("<b>Estadísticas</b>", styles['Heading2'])
    elements.append(stats_title)
    elements.append(Spacer(1, 0.1*inch))
    
    stats = [
        f'<b>Total de Citas:</b> {len(citas)}',
        f'<b>Confirmadas:</b> {sum(1 for c in citas if c.get("estado") == "confirmada")}',
        f'<b>Pendientes:</b> {sum(1 for c in citas if c.get("estado") == "pendiente")}',
        f'<b>Canceladas:</b> {sum(1 for c in citas if c.get("estado") == "cancelada")}',
    ]
    
    for stat in stats:
        elements.append(Paragraph(stat, styles['Normal']))
    
    doc.build(elements)
    output.seek(0)
    return output


def exportar_atenciones_a_pdf(atenciones):
    """Exporta atenciones a PDF"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    elements = []
    
    # Estilos
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#70AD47'),
        spaceAfter=30,
        alignment=1  # CENTER
    )
    
    # Título
    title = Paragraph("Reporte de Atenciones", title_style)
    elements.append(title)
    
    # Fecha
    fecha_text = Paragraph(f'<b>Generado:</b> {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}', styles['Normal'])
    elements.append(fecha_text)
    elements.append(Spacer(1, 0.3*inch))
    
    # Tabla
    data = [["ID", "Paciente", "Odontólogo", "Fecha Inicio", "Fecha Fin", "Estado"]]
    for atencion in atenciones:
        paciente = atencion.get('id_paciente', {})
        if isinstance(paciente, dict):
            paciente = paciente.get('nombre', '')
        
        odontologo = atencion.get('id_odontologo', {})
        if isinstance(odontologo, dict):
            odontologo = odontologo.get('nombre', '')
        
        data.append([
            str(atencion.get('id_atencion', '')),
            str(paciente),
            str(odontologo),
            str(atencion.get('fecha_inicio', '')),
            str(atencion.get('fecha_fin', '')),
            str(atencion.get('estado', ''))
        ])
    
    table = Table(data, colWidths=[0.6*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#70AD47')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Estadísticas
    stats_title = Paragraph("<b>Estadísticas</b>", styles['Heading2'])
    elements.append(stats_title)
    elements.append(Spacer(1, 0.1*inch))
    
    stats = [
        f'<b>Total de Atenciones:</b> {len(atenciones)}',
        f'<b>Finalizadas:</b> {sum(1 for a in atenciones if a.get("estado") == "finalizada")}',
        f'<b>En Curso:</b> {sum(1 for a in atenciones if a.get("estado") == "en_curso")}',
        f'<b>Canceladas:</b> {sum(1 for a in atenciones if a.get("estado") == "cancelada")}',
    ]
    
    for stat in stats:
        elements.append(Paragraph(stat, styles['Normal']))
    
    doc.build(elements)
    output.seek(0)
    return output
